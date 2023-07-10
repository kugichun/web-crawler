import docx


if __name__ == '__main__':
    # 获取文档对象
    doc_path = 'result.docx'
    doc = docx.Document(doc_path)
    print("段落数:" + str(len(doc.paragraphs)))  # 段落数

    # 输出每一段的内容
    for para in doc.paragraphs:
        print(para.text)
