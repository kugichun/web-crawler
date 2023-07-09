import requests
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup
import docx
from docx import Document
from docx.shared import Inches
import re

b=""
#########################################################################
# 定義函式，以取得html_page文本


def get_webpage(url):
    html_page = requests.get(url)

    if html_page.status_code != 200:
        print("invalid url", html_page.status_code)
        return None
    else:
        return html_page.text

#########################################################################
# 嘗試抓特定區塊資料


site = "https://www.informationsecurity.com.tw/article/json/ajax_list.aspx?mod=1"
html = get_webpage(site)
soup = BeautifulSoup(html, "html.parser")


# 標題+連結
http = soup.find_all("a", class_="articleList_box")
for link in http:
    title = link.text
    news_link = 'https://www.informationsecurity.com.tw'+link['href']

    
    # 抓內文資料
    site_in = "https://www.informationsecurity.com.tw"+str(link['href'])
    html_in = get_webpage(site_in)
    soup_in = BeautifulSoup(html_in, "html.parser")
    keyword = soup_in.find_all("div", class_="tag_box w_93")
    for key in keyword:
        #print(key.text)
        # 判斷關鍵字是否符合
        if ("網路釣魚" in key.text or "勒索軟體攻擊" in key.text):
            '''
            print("----------------------------------------------------------------")
            print("抓到") 
            print(title)      # 標題
            print(news_link)  # 連結
            print("----------------------------------------------------------------")
            '''
            a=title+"\n"+news_link+"------------------------------------------------------------------------------------------------"
            b=b+a
print(b) 
                 
################################################
        #else:
        #print("不符")
        
#############################################################
# 符合輸出成word
# Initialise the Word document
    
doc = docx.Document()
doc.add_heading("Informationsecurity News", level=0)
new=doc.add_paragraph(b)
                
# Save the Word doc
doc.save('result.docx')
            
        

########################################################################