#完整版
import requests
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup
import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import re
################################################################
#收集多個瀏覽器的User-Agent，每次發起請求時隨機抽取一個使用，可以進一步提高安全性，並避免被封IP
import requests,random
user_agents = ['User-Agent:Mozilla/5.0 (Windows NT 6.1; rv:2.0.1) Gecko/20100101 Firefox/4.0.1','User-Agent:Mozilla/5.0 (Windows; U; Windows NT 6.1; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50','User-Agent:Opera/9.80 (Windows NT 6.1; U; en) Presto/2.8.131 Version/11.11']
def get_html(url):
   headers = {'User-Agent':random.choice(user_agents)}
   resp = requests.get(url,headers = headers)
   return resp.text
###############################################################

#########################################################################
# 定義函式，以取得html_page文本

def get_webpage(url):
    html_page = requests.get(url)

    if html_page.status_code != 200:
        print("invalid url", html_page.status_code)
        return None
    else:
        return html_page.text

#########################################################################
# 嘗試抓特定區塊資料

site = "https://www.informationsecurity.com.tw/article/json/ajax_list.aspx?mod=1"
html = get_webpage(site)
soup = BeautifulSoup(html, "html.parser")

# 標題+連結
#http = soup.find_all("a", class_="articleList_box")
http2 = soup.find("a", class_="articleList_box")
news_link_0 = 'https://www.informationsecurity.com.tw'+http2['href'] #抓最新的那則新聞，目的是取出當中的最大index

MaxIndex = re.findall(r"\d+", news_link_0)[0]
MaxIndex = int(MaxIndex)
initial = 10500
output="" 
for i in range(0,MaxIndex-initial):

    #title = http.text
    news_link = 'https://www.informationsecurity.com.tw/article/article_detail.aspx?aid='+str(MaxIndex-i)+'&mod=1'

    # 抓內文資料-標題
    site_in = news_link
    html_in = get_webpage(site_in)
    soup_in = BeautifulSoup(html_in, "html.parser")
    #title = soup_in.find("h1", class_="title_content w_93").text
    #print(title)

    """
    # 抓內文資料-關鍵字
    site_in = news_link
    html_in = get_webpage(site_in)
    soup_in = BeautifulSoup(html_in, "html.parser")
    """

    #抓關鍵字
    keyword = soup_in.find_all("div", class_="tag_box w_93")
    for key in keyword:
        # 判斷關鍵字是否符合
        if ("供應鏈安全" in key.text or "勒索軟體攻擊" in key.text):
            '''
            print("----------------------------------------------------------------")
            print("抓到") 
            print(title)      # 標題
            print(news_link)  # 連結
            print("----------------------------------------------------------------")
            '''
            title = soup_in.find("h1", class_="title_content w_93").text
            news=title+"\n"+news_link+"\n------------------------------------------------------------------------------------------------\n"
            output=output+news
            print(news) 

#print(output) 

initial=MaxIndex #更新起點
            
################################################
        #else:
        #print("不符")       
#############################################################
"""
# 第一次創建新word檔請用這段(請注意，會覆蓋，僅限第一次創建)
# 輸出成word
 
doc = docx.Document()
doc.add_heading("Informationsecurity News", level=0)
new=doc.add_paragraph(output)
                
# Save the Word doc
doc.save('result.docx')
"""
########################################################################
# 第二次後更新word檔資料(請注意，須先具有第一次創建的檔案)
# 輸出成word

#取的原本的內文(目的是使新資料能夠insert到最前面，需要定位)
if __name__ == '__main__':
    # 獲取文檔
    doc = docx.Document('result.docx')
    # 輸出內容
    data=''
    for para in doc.paragraphs:
        data=data+para.text
#print(data)

doc = docx.Document()
doc.add_heading("Informationsecurity News", level=0)
doc.add_heading("近期更新", level=1)
news=doc.add_paragraph(output)
doc.add_heading("過去資料", level=1)
paragraph=doc.add_paragraph(data)

# Save the Word doc
doc.save('result.docx')

########################################################################