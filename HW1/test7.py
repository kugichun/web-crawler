###爬蟲完整版+檢測結果###

import requests
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup
import docx
from docx import Document
from docx.shared import Inches
import re
import statistics

####################################################
# 定義函式，以取得html_page文本


def get_webpage(url):
    html_page = requests.get(url)

    if html_page.status_code != 200:
        print("invalid url", html_page.status_code)
        return None
    else:
        return html_page.text

###############################################################################
# 取得總頁數


site = "https://www.tenable.com/plugins/nessus/families/Backdoors?page"
html = get_webpage(site)
soup = BeautifulSoup(html, "html.parser")
total_pages = soup.find_all(
    "a", class_="page-link page-text")[0]

num = re.findall(r"\d+", total_pages.text)
num = np.array(num)
print(num[1])

total_pages = int(num[1])
################################################
# 讀取各頁資料
df = pd.DataFrame()
for page in range(1, total_pages+1):
    site = "https://www.tenable.com/plugins/nessus/families/Backdoors?page=" + \
        str(page)
    html = get_webpage(site)

    df = pd.concat([df, pd.read_html(html)[0]])
#df = df1[0]
#df.columns = ["ID", "Name", "Severity"]
df["Empty"] = ""  # 新增空欄


print(df)
print("型態：", type(df))
print("長度：", len(df))
# tmp += df
# print(df)
# 3

########################################################################
# Initialise the Word document
doc = docx.Document()

doc.add_heading("Backdoors Family for Nessus", level=0)

# 標題
table = doc.add_table(rows=1, cols=4, style='Table Grid')
hc = table.rows[0].cells
hc[0].text = 'ID'
hc[1].text = 'Name'
hc[2].text = 'Severity'
hc[3].text = '檢測結果'

# Initialise the table
t = doc.add_table(rows=df.shape[0], cols=df.shape[1])
t.style = 'Table Grid'

# Add the body of the data frame to the table
for i in range(df.shape[0]):
    for j in range(df.shape[1]):
        cell = df.iat[i, j]
        t.cell(i, j).text = str(cell)

# Save the Word doc
doc.save('table2.docx')
